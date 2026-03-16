"""
================================================================================
                    DOCUMENT SERVICE — Create MD, DOCX, PDF
================================================================================
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Optional
import logging
import markdown
from docx import Document
from docx.shared import Inches, Pt
from weasyprint import HTML, CSS

logger = logging.getLogger("DOCUMENT_SERVICE")

ARTEFACTS_ROOT = Path(os.getenv("ARTEFACTS_ROOT", "/data/artefacts"))


class DocumentService:
    """Service to create documents from LLM content"""
    
    def __init__(self):
        self.artefacts_root = ARTEFACTS_ROOT
        self.artefacts_root.mkdir(parents=True, exist_ok=True)
    
    def create_markdown(self, content: str, filename: str) -> Path:
        """Create a markdown file"""
        if not filename.endswith('.md'):
            filename += '.md'
        
        filepath = self.artefacts_root / filename
        filepath.write_text(content, encoding='utf-8')
        logger.info(f"Created MD: {filepath}")
        return filepath
    
    def create_docx(self, content: str, filename: str, title: Optional[str] = None) -> Path:
        """Create a DOCX file from markdown content"""
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        doc = Document()
        
        # Add title if provided
        if title:
            doc.add_heading(title, 0)
        
        # Parse markdown and add to doc
        lines = content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                doc.add_paragraph(line[3:], style='List Number')
            elif line.strip() == '':
                continue
            else:
                doc.add_paragraph(line)
        
        filepath = self.artefacts_root / filename
        doc.save(filepath)
        logger.info(f"Created DOCX: {filepath}")
        return filepath
    
    def create_pdf(self, content: str, filename: str, title: Optional[str] = None) -> Path:
        """Create a PDF file from markdown content"""
        if not filename.endswith('.pdf'):
            filename += '.pdf'
        
        # Convert markdown to HTML
        html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        
        # Wrap in full HTML with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                h2 {{ color: #555; }}
                h3 {{ color: #666; }}
                code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background: #f4f4f4; }}
            </style>
        </head>
        <body>
            {f'<h1>{title}</h1>' if title else ''}
            {html_content}
        </body>
        </html>
        """
        
        filepath = self.artefacts_root / filename
        HTML(string=full_html).write_pdf(filepath)
        logger.info(f"Created PDF: {filepath}")
        return filepath
    
    def update_document(self, filepath: Path, new_content: str) -> Path:
        """Update an existing document with new content"""
        ext = filepath.suffix.lower()
        filename = filepath.name
        
        if ext == '.md':
            return self.create_markdown(new_content, filename)
        elif ext == '.docx':
            return self.create_docx(new_content, filename)
        elif ext == '.pdf':
            return self.create_pdf(new_content, filename)
        else:
            raise ValueError(f"Unsupported format: {ext}")


# Singleton
document_service = DocumentService()
